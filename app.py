import streamlit as st
import pandas as pd
import requests
from urllib.parse import urlparse
from concurrent.futures import ThreadPoolExecutor, as_completed
import xml.etree.ElementTree as ET
from io import BytesIO

st.set_page_config(page_title="IA Audit Tool", layout="wide")
st.title("🚀 IA Audit Tool (Advanced)")

# -------------------------------
# Safe Imports
# -------------------------------
BS4_AVAILABLE, DOCX_AVAILABLE, EXCEL_AVAILABLE = True, True, True

try:
    from bs4 import BeautifulSoup
except:
    BS4_AVAILABLE = False

try:
    from docx import Document
except:
    DOCX_AVAILABLE = False

try:
    import openpyxl
except:
    EXCEL_AVAILABLE = False

HEADERS = {"User-Agent": "Mozilla/5.0"}

# -------------------------------
# Sidebar
# -------------------------------
st.sidebar.header("⚙️ Settings")
threads = st.sidebar.slider("Parallel Threads", 5, 30, 15)

# -------------------------------
# Inputs
# -------------------------------
uploaded_file = st.file_uploader("Upload URL file", type=["csv", "xlsx"])
sitemap_url = st.text_input("Optional Sitemap URL (for orphan detection)")

# -------------------------------
# Helpers
# -------------------------------
def load_file(file):
    return pd.read_csv(file) if file.name.endswith(".csv") else pd.read_excel(file)

def normalize(df):
    df.columns = [str(c).lower() for c in df.columns]
    for col in df.columns:
        if col in ["url", "link"]:
            return df[[col]].rename(columns={col: "URL"})
    return df.iloc[:, [0]].rename(columns={df.columns[0]: "URL"})

def enrich(df):
    df["path"] = df["URL"].apply(lambda x: urlparse(x).path)
    df["depth"] = df["path"].apply(lambda x: len([p for p in x.split("/") if p]))
    df["section"] = df["path"].apply(
        lambda x: x.split("/")[1] if len(x.split("/")) > 1 else "root"
    )
    return df

def fetch(url):
    try:
        r = requests.get(url, timeout=5, headers=HEADERS)
        title = ""
        if BS4_AVAILABLE:
            soup = BeautifulSoup(r.text, "lxml")
            title = soup.title.string.strip() if soup.title else ""
        return r.status_code, r.url, title
    except:
        return None, None, ""

def crawl(df):
    urls = df["URL"].tolist()
    status, final, title = [None]*len(urls), [None]*len(urls), [""]*len(urls)

    progress = st.progress(0)

    with ThreadPoolExecutor(max_workers=threads) as exe:
        futures = {exe.submit(fetch, u): i for i, u in enumerate(urls)}

        for i, f in enumerate(as_completed(futures)):
            idx = futures[f]
            s, fu, t = f.result()
            status[idx], final[idx], title[idx] = s, fu, t
            progress.progress((i+1)/len(urls))

    df["status"], df["final_url"], df["title"] = status, final, title
    return df

def get_sitemap(url):
    try:
        r = requests.get(url)
        root = ET.fromstring(r.content)
        return set([e.text for e in root.iter() if "loc" in e.tag])
    except:
        return set()

# -------------------------------
# Metrics
# -------------------------------
def compute_metrics(df, sitemap):
    total = len(df)
    unique = df["URL"].nunique()
    duplicates = total - unique

    dup_titles_df = df[df["title"].duplicated(keep=False)]

    section_counts = df["section"].value_counts()
    section_pct = (section_counts / total * 100).round(2)

    deep_pct = round((df[df["depth"] > 4].shape[0] / total) * 100, 2)

    # Orphan pages
    if sitemap:
        orphan_urls = list(set(sitemap) - set(df["URL"]))
        orphan_pct = round((len(orphan_urls) / len(sitemap)) * 100, 2)
    else:
        orphan_urls, orphan_pct = [], "N/A"

    metrics = {
        "Total Pages": total,
        "% Duplicate URLs": round((duplicates/total)*100,2) if total else 0,
        "% Duplicate Titles": round((len(dup_titles_df)/total)*100,2),
        "Avg Depth": round(df["depth"].mean(),2),
        "% Pages > Depth 4": deep_pct,
        "Broken Pages": df[df["status"]>=400].shape[0],
        "Redirects": df[df["status"].between(300,399)].shape[0],
        "% Orphan Pages": orphan_pct
    }

    return metrics, section_counts, section_pct, dup_titles_df, orphan_urls

# -------------------------------
# Narrative Summary
# -------------------------------
def generate_summary(metrics, section_pct):
    return f"""
### IA Audit Summary

- Total Pages: {metrics['Total Pages']}
- Duplicate Content: ~{metrics['% Duplicate URLs']}%
- Average Depth: {metrics['Avg Depth']}
- Deep Pages (>4 levels): {metrics['% Pages > Depth 4']}%

### Navigation Health
- Orphan Pages: {metrics['% Orphan Pages']}%
- Broken Pages: {metrics['Broken Pages']}
- Redirects: {metrics['Redirects']}

### Content Distribution (%)
{section_pct.to_string()}

### Key Observations
- High duplication indicates structural inefficiencies
- Deep navigation reduces discoverability
- Fragmented entry points likely exist

### Recommendation
- Move to entity-driven architecture
- Reduce duplication via centralized content
- Improve navigation depth and structure
"""

# -------------------------------
# Reports
# -------------------------------
def build_excel(df, metrics, sections):
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, "Raw Data", index=False)
        pd.DataFrame(metrics.items(), columns=["Metric","Value"]).to_excel(writer, "Metrics", index=False)
        sections.to_frame("Count").to_excel(writer, "Sections")
    buffer.seek(0)
    return buffer

def build_word(metrics, sections, summary):
    doc = Document()
    doc.add_heading("IA Audit Report", 0)

    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(summary)

    doc.add_heading("Metrics", 1)
    for k,v in metrics.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Section Distribution", 1)
    for s,c in sections.items():
        doc.add_paragraph(f"{s}: {c}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -------------------------------
# MAIN FLOW
# -------------------------------
if uploaded_file:
    df = normalize(load_file(uploaded_file)).dropna()
    df = enrich(df)

    st.info("⚡ Running parallel crawl...")
    df = crawl(df)

    sitemap = get_sitemap(sitemap_url) if sitemap_url else set()

    metrics, sections, section_pct, dup_titles_df, orphan_urls = compute_metrics(df, sitemap)
    summary = generate_summary(metrics, section_pct)

    # -------------------------------
    # UI Tabs
    # -------------------------------
    tab1, tab2, tab3, tab4 = st.tabs(["📊 Dashboard","🔍 Insights","📁 Data","⬇️ Reports"])

    with tab1:
        cols = st.columns(len(metrics))
        for i,(k,v) in enumerate(metrics.items()):
            cols[i].metric(k,v)

        st.bar_chart(section_pct)

    with tab2:
        st.markdown(summary)

        st.subheader("Duplicate Pages (Title Level)")
        st.dataframe(dup_titles_df[["URL","title"]].head(50))

        if orphan_urls:
            st.subheader("Orphan Pages")
            st.dataframe(pd.DataFrame(orphan_urls, columns=["URL"]).head(50))

    with tab3:
        st.dataframe(df)

    with tab4:
        if EXCEL_AVAILABLE:
            st.download_button("Download Excel Report", build_excel(df, metrics, sections), "IA_Report.xlsx")

        if DOCX_AVAILABLE:
            st.download_button("Download Word Report", build_word(metrics, sections, summary), "IA_Report.docx")